# -*- coding: utf-8 -*-
"""Génère la documentation Word 'Styles & properties' pour clients non techniques."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x2F, 0x53, 0xC8)
INK = RGBColor(0x17, 0x20, 0x2B)
MUTED = RGBColor(0x56, 0x64, 0x6F)
SCREEN = RGBColor(0x9A, 0x6A, 0x15)

doc = Document()

# Styles de base
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = INK

def spacer(size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(size)
    return p

def screenshot(note):
    p = doc.add_paragraph()
    r = p.add_run('\U0001F4F8  Screenshot à prendre : ' + note)
    r.bold = True
    r.font.color.rgb = SCREEN
    r.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)

def tip(text):
    p = doc.add_paragraph()
    r = p.add_run('\U0001F4A1  ' + text)
    r.italic = True
    r.font.color.rgb = MUTED
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(6)

def warn(text):
    p = doc.add_paragraph()
    r = p.add_run('⚠️  ' + text)
    r.italic = True
    r.font.color.rgb = SCREEN
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(6)

def options_table(rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, label in enumerate(['Option', 'En clair', "Quand l'utiliser"]):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(10.5)
    for opt, clear, use in rows:
        cells = table.add_row().cells
        cells[0].text = ''
        r = cells[0].paragraphs[0].add_run(opt)
        r.bold = True
        r.font.size = Pt(10.5)
        for idx, txt in [(1, clear), (2, use)]:
            cells[idx].text = ''
            rr = cells[idx].paragraphs[0].add_run(txt)
            rr.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ---------- Titre ----------
title = doc.add_heading('Styles & properties — la section « Général »', level=0)
for run in title.runs:
    run.font.color.rgb = INK

intro = doc.add_paragraph()
ri = intro.add_run(
    "Ces réglages servent à affiner la mise en page d'un élément. "
    "Dans la plupart des cas vous n'en aurez pas besoin — mais quand vous voulez un rendu "
    "précis (mettre des éléments côte à côte, coller un bouton en bas de l'écran, "
    "habiller une image de texte…), c'est ici que ça se passe."
)
ri.font.color.rgb = MUTED
intro.paragraph_format.space_after = Pt(8)

screenshot("la section GÉNÉRAL entière du panneau de droite (Affichage + Alignement flottant + Position).")

# ---------- 1. Affichage ----------
doc.add_heading('1. Affichage — comment l’élément occupe la place', level=1)
screenshot("le menu déroulant Affichage ouvert (block / inline / inline-block / flex / none).")
options_table([
    ('block', "L'élément prend toute la largeur et passe à la ligne, comme un paragraphe.",
     "Le réglage normal d'un titre, d'un paragraphe, d'une section."),
    ('inline', "L'élément se met dans le fil du texte, collé aux autres, comme un mot dans une phrase. On ne peut pas régler sa largeur/hauteur.",
     "Un mot ou un lien à l'intérieur d'un texte."),
    ('inline-block', "Comme inline (les éléments se mettent côte à côte) mais on peut régler sa taille et ses marges.",
     "Plusieurs petits boutons ou étiquettes alignés côte à côte."),
    ('flex', "Transforme l'élément en boîte « intelligente » qui range ses éléments côte à côte et permet de les aligner/espacer facilement.",
     "Aligner proprement plusieurs blocs sur une ligne (3 colonnes, une rangée d'icônes)."),
    ('none', "Cache complètement l'élément (invisible, ne prend plus de place).",
     "Masquer temporairement un bloc sans le supprimer."),
])
tip("Image mentale : block = une ligne à soi tout seul ; inline / inline-block = des mots côte à côte ; flex = une étagère qui range ses objets en ligne.")

# ---------- 2. Alignement flottant ----------
doc.add_heading('2. Alignement flottant — faire « flotter » un élément et l’habiller de texte', level=1)
screenshot("la ligne ALIGNEMENT FLOTTANT (none / left / right).")
options_table([
    ('none', "Pas de flottement (réglage normal).", "Par défaut."),
    ('left', "L'élément se colle à gauche, et le texte vient l'entourer à droite.",
     "Une image à gauche avec le texte qui s'enroule autour."),
    ('right', "L'élément se colle à droite, le texte l'entoure à gauche.",
     "Une image à droite habillée de texte."),
])
tip("Exemple concret : une petite photo à gauche d'un paragraphe, avec le texte qui coule autour — comme dans un article de magazine.")
warn("Conseil : pour aligner des blocs entiers côte à côte, préférez Affichage → flex (plus simple et plus fiable). Le flottant sert surtout à habiller une image de texte.")

# ---------- 3. Position ----------
doc.add_heading('3. Position — où et comment l’élément est placé', level=1)
screenshot("la ligne POSITION (static / relative / absolute / fixed).")
options_table([
    ('static', "Position normale, l'élément reste dans le flux de la page.", "Par défaut, la plupart du temps."),
    ('relative', "On peut décaler légèrement l'élément par rapport à sa place normale, sans déranger les autres.",
     "Petit ajustement ; sert aussi de repère pour un élément en absolute."),
    ('absolute', "L'élément se place librement par-dessus le reste, à l'endroit qu'on veut.",
     "Superposer : un badge « Nouveau » sur une image, une étiquette sur une photo."),
    ('fixed', "L'élément reste collé à l'écran même quand on fait défiler la page.",
     "Un bandeau ou un bouton qui suit (barre « S'inscrire » toujours visible en bas)."),
])
tip("Image mentale : static/relative = l'élément reste à sa place ; absolute = un autocollant qu'on pose où on veut sur la page ; fixed = un autocollant collé sur l'écran, qui ne bouge pas quand on défile.")

# ---------- Récap screenshots ----------
doc.add_heading('\U0001F4F8 Récap des screenshots à prendre', level=1)
for t in [
    "La section GÉNÉRAL complète (vue d'ensemble). — déjà pris",
    "Le menu Affichage ouvert (les 5 options). — déjà pris",
    "La ligne Alignement flottant (none / left / right).",
    "La ligne Position (static / relative / absolute / fixed).",
]:
    doc.add_paragraph(t, style='List Number')

p = doc.add_paragraph()
r = p.add_run("Optionnel mais très parlant pour un client non technique — 2 ou 3 exemples « avant / après » dans le canevas :")
r.bold = True
for t in [
    "Flex : 3 blocs empilés (avant) → les 3 côte à côte (après).",
    "Flottant left : une image + du texte qui l'entoure.",
    "Position fixed : un bouton qui reste visible en bas quand on fait défiler.",
]:
    doc.add_paragraph(t, style='List Bullet')

out = r"C:\Users\ZahiraElmelsse\Desktop\monPFE\LandingPageGenerator\Documentation_Styles_Properties.docx"
doc.save(out)
print("SAVED:", out)
